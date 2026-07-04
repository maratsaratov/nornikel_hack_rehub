"""Smoke checks for the standalone ingestion subsystem.

Run from the repository root:
    python backend/smoke_ingestion.py
"""

import importlib.util
import os
import tempfile
from io import BytesIO


def _has_module(name):
    return importlib.util.find_spec(name) is not None


def _sample_pdf_bytes():
    stream = b"BT /F1 12 Tf 72 720 Td (Hypothesis factory parser smoke PDF 2024) Tj ET"
    objects = [
        b"1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj\n",
        b"2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj\n",
        (
            b"3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >> endobj\n"
        ),
        b"4 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj\n",
        b"5 0 obj << /Length " + str(len(stream)).encode("ascii") + b" >> stream\n"
        + stream + b"\nendstream endobj\n",
    ]
    body = b"%PDF-1.4\n"
    offsets = [0]
    for obj in objects:
        offsets.append(len(body))
        body += obj
    xref_start = len(body)
    body += f"xref\n0 {len(objects) + 1}\n".encode("ascii")
    body += b"0000000000 65535 f \n"
    for offset in offsets[1:]:
        body += f"{offset:010d} 00000 n \n".encode("ascii")
    body += (
        b"trailer << /Size " + str(len(objects) + 1).encode("ascii") +
        b" /Root 1 0 R >>\nstartxref\n" + str(xref_start).encode("ascii") +
        b"\n%%EOF\n"
    )
    return body


def _sample_xlsx_bytes():
    from openpyxl import Workbook

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "experiments"
    sheet.append(["experiment", "temperature", "yield"])
    sheet.append(["E-1", 720, 0.82])
    sheet.append(["E-2", 760, 0.87])
    stream = BytesIO()
    workbook.save(stream)
    return stream.getvalue()


def _sample_docx_bytes():
    from docx import Document

    document = Document()
    document.add_heading("Parser smoke report", level=1)
    document.add_paragraph("The document contains deterministic extraction content for 2024.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "parameter"
    table.cell(0, 1).text = "value"
    table.cell(1, 0).text = "temperature"
    table.cell(1, 1).text = "720"
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def _samples():
    samples = [
        ("txt", "sample.txt", b"Parser Smoke\n\nThis text document validates chunk creation for 2024."),
        ("csv", "sample.csv", b"experiment,temperature,yield\nE-1,720,0.82\nE-2,760,0.87\n"),
    ]
    if _has_module("openpyxl"):
        samples.append(("xlsx", "sample.xlsx", _sample_xlsx_bytes()))
    if _has_module("docx"):
        samples.append(("docx", "sample.docx", _sample_docx_bytes()))
    if _has_module("pypdf"):
        samples.append(("pdf", "sample.pdf", _sample_pdf_bytes()))
    return samples


def run_api_smoke():
    temp_dir = tempfile.mkdtemp(prefix="ingestion-smoke-")
    os.environ["DATABASE_URL"] = "sqlite:///:memory:"
    os.environ["SEED_DEMO"] = "false"
    os.environ["UPLOAD_DIR"] = os.path.join(temp_dir, "uploads")

    from fastapi.testclient import TestClient
    from app import app
    from db import db
    from models import DocumentChunk, DocumentTable, Project, SourceDocument

    with TestClient(app) as client:
        project = Project(title="Smoke project", kpi_target="Increase experimental yield")
        db.session.add(project)
        db.session.commit()
        project_id = project.id

        for file_type, filename, payload in _samples():
            response = client.post(
                f"/api/projects/{project_id}/documents?parse=true",
                files={"file": (filename, payload, "application/octet-stream")},
            )
            assert response.status_code == 201, response.text
            data = response.json()
            document = data["document"]
            assert document["parse_status"] == "parsed", data
            assert document["file_type"] == file_type, data

        assert SourceDocument.query.count() >= 2
        assert DocumentChunk.query.count() >= 2
        assert DocumentTable.query.count() >= 1

    print("API ingestion smoke passed")


def run_extractor_smoke():
    from ingestion.chunking import build_chunks
    from ingestion.extractors.csv_extractor import CSVExtractor
    from ingestion.extractors.docx_extractor import DOCXExtractor
    from ingestion.extractors.pdf_extractor import PDFExtractor
    from ingestion.extractors.txt_extractor import TXTExtractor
    from ingestion.extractors.xlsx_extractor import XLSXExtractor

    extractors = {
        "txt": TXTExtractor,
        "csv": CSVExtractor,
        "xlsx": XLSXExtractor,
        "docx": DOCXExtractor,
        "pdf": PDFExtractor,
    }
    with tempfile.TemporaryDirectory(prefix="ingestion-extractors-") as temp_dir:
        for file_type, filename, payload in _samples():
            path = os.path.join(temp_dir, filename)
            with open(path, "wb") as handle:
                handle.write(payload)
            result = extractors[file_type]().extract(path, filename)
            chunks = build_chunks(result.sections, result.tables)
            assert result.file_type == file_type
            assert result.raw_text or result.tables or result.warnings
            assert chunks or file_type == "pdf"
            print(f"{file_type}: sections={len(result.sections)} tables={len(result.tables)} chunks={len(chunks)}")
    print("Extractor ingestion smoke passed")


if __name__ == "__main__":
    web_modules = ("fastapi", "sqlalchemy", "openai", "sklearn")
    if all(_has_module(name) for name in web_modules):
        run_api_smoke()
    else:
        print("Project web dependencies are not installed; running extractor-only smoke.")
        run_extractor_smoke()


"""Экспорт данных проекта в разные форматы.

  • PDF / DOCX — бизнес-отчёт: ранжированные гипотезы + достижение/прогноз KPI;
  • CSV / JSON — «форматы задач»: плоский/структурный список гипотез для трекеров;
  • Markdown  — вспомогательный текстовый формат.

reportlab (PDF) и python-docx (DOCX) импортируются лениво — модуль работает даже
без них (CSV/JSON/MD не зависят от тяжёлых библиотек). Для кириллицы в PDF ищем
Unicode-TTF (DejaVu в Docker, Arial/Segoe UI локально, либо путь из EXPORT_FONT_*).
"""
import csv
import io
import json
import os
from collections import Counter
from datetime import datetime
from xml.sax.saxutils import escape as _xml_escape


class ExportError(RuntimeError):
    """Формат недоступен (например, нет шрифта для PDF)."""


STATUS_LABELS = {
    "proposed": "Новая",
    "review": "На проверке",
    "accepted": "Подтверждена экспертом",
    "rejected": "Опровергнута экспертом",
}
DIMS = [("novelty", "Новизна"), ("value", "Ценность / эффект"),
        ("feasibility", "Реализуемость"), ("risk", "Риск")]


def _scores(h):
    return h.get("effective_scores") or h.get("scores") or {}


def _num(v):
    try:
        return round(float(v))
    except (TypeError, ValueError):
        return 0


def _summary(hyps):
    total = len(hyps)
    by_status = Counter(h.get("status", "proposed") for h in hyps)
    avg = round(sum(float(h.get("composite") or 0) for h in hyps) / total, 1) if total else 0
    high_value = sum(1 for h in hyps if float(_scores(h).get("value") or 0) >= 70)
    top = hyps[0] if hyps else None
    return {
        "total": total,
        "confirmed": by_status.get("accepted", 0),
        "refuted": by_status.get("rejected", 0),
        "review": by_status.get("review", 0),
        "proposed": by_status.get("proposed", 0),
        "avg_composite": avg,
        "high_value_count": high_value,
        "top_statement": top.get("statement") if top else None,
        "top_composite": top.get("composite") if top else None,
    }


def _kpi_line(project):
    bits = []
    if project.get("kpi_metric"):
        direction = "увеличить" if (project.get("kpi_direction") or "increase") == "increase" else "снизить"
        bits.append(f"Метрика: {project['kpi_metric']} ({direction})")
    if project.get("domain"):
        bits.append(f"Область: {project['domain']}")
    return " · ".join(bits)


# ── JSON / CSV / Markdown («форматы задач» + текст) ──────────────────────────
def to_json(project, hyps, weights):
    payload = {
        "exported_at": datetime.utcnow().isoformat() + "Z",
        "project": project,
        "ranking_weights": weights,
        "summary": _summary(hyps),
        "hypotheses": [{
            "rank": i,
            "composite": h.get("composite"),
            "status": h.get("status"),
            "statement": h.get("statement"),
            "goal_link": h.get("goal_link"),
            "mechanism": h.get("mechanism"),
            "rationale": h.get("rationale"),
            "validation": h.get("validation"),
            "scores": _scores(h),
            "model_scores": h.get("scores"),
            "expert_notes": h.get("expert_notes"),
            "tags": h.get("tags") or [],
            "sources": [e.get("title") for e in (h.get("evidence") or []) if e.get("title")],
        } for i, h in enumerate(hyps, 1)],
    }
    return json.dumps(payload, ensure_ascii=False, indent=2)


def to_csv(hyps):
    buf = io.StringIO()
    w = csv.writer(buf)
    w.writerow([
        "rank", "composite", "status", "statement", "goal_link", "mechanism", "validation",
        "novelty", "value", "feasibility", "risk", "expert_notes", "tags", "sources",
    ])
    for i, h in enumerate(hyps, 1):
        s = _scores(h)
        srcs = "; ".join(filter(None, [e.get("title") for e in (h.get("evidence") or [])]))
        w.writerow([
            i, h.get("composite"), h.get("status"), h.get("statement"), h.get("goal_link") or "",
            h.get("mechanism") or "", h.get("validation") or "",
            s.get("novelty"), s.get("value"), s.get("feasibility"), s.get("risk"),
            h.get("expert_notes") or "", ", ".join(h.get("tags") or []), srcs,
        ])
    return "﻿" + buf.getvalue()  # BOM — корректная кириллица в Excel


def to_markdown(project, hyps, weights):
    L = [f"# Гипотезы НИОКР — {project.get('title', '')}\n",
         f"**Цель (KPI):** {project.get('kpi_target', '')}  "]
    if _kpi_line(project):
        L.append(f"**{_kpi_line(project)}**  ")
    L.append(f"\nВсего гипотез: {len(hyps)}\n\n---\n")
    for i, h in enumerate(hyps, 1):
        s = _scores(h)
        L.append(f"## {i}. {h.get('statement', '')}\n")
        L.append(f"**Ранг:** {h.get('composite')} · **Статус:** {STATUS_LABELS.get(h.get('status'), h.get('status'))}\n")
        for key, val in (("goal_link", "Связь с целью"), ("mechanism", "Механизм"),
                         ("validation", "План проверки"), ("expert_notes", "Отзыв эксперта")):
            if h.get(key):
                L.append(f"**{val}:** {h[key]}\n")
        L.append(f"**Оценки:** новизна {s.get('novelty')}, ценность {s.get('value')}, "
                 f"реализуемость {s.get('feasibility')}, риск {s.get('risk')}\n\n---\n")
    return "\n".join(L)


# ── Шрифт для PDF ────────────────────────────────────────────────────────────
_FONT_CANDIDATES = [
    (os.getenv("EXPORT_FONT_REGULAR"), os.getenv("EXPORT_FONT_BOLD")),
    ("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
    ("C:/Windows/Fonts/arial.ttf", "C:/Windows/Fonts/arialbd.ttf"),
    ("C:/Windows/Fonts/segoeui.ttf", "C:/Windows/Fonts/segoeuib.ttf"),
    ("/Library/Fonts/Arial.ttf", "/Library/Fonts/Arial Bold.ttf"),
]
_FONT, _FONT_B = "HFBody", "HFBold"
_fonts_ready = False


def _ensure_pdf_fonts():
    global _fonts_ready
    if _fonts_ready:
        return
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    reg = bold = None
    for cand_reg, cand_bold in _FONT_CANDIDATES:
        if cand_reg and os.path.exists(cand_reg):
            reg = cand_reg
            bold = cand_bold if (cand_bold and os.path.exists(cand_bold)) else cand_reg
            break
    if not reg:
        raise ExportError(
            "Для PDF нужен Unicode-TTF с кириллицей. Установите fonts-dejavu-core "
            "(в Docker уже включён) или задайте EXPORT_FONT_REGULAR. Пока используйте DOCX."
        )
    pdfmetrics.registerFont(TTFont(_FONT, reg))
    pdfmetrics.registerFont(TTFont(_FONT_B, bold))
    _fonts_ready = True


def _p(text):
    return _xml_escape(str(text or ""))


def to_pdf(project, hyps, weights):
    _ensure_pdf_fonts()
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

    def st(name, font, size, lead, color, sb=0, sa=3):
        return ParagraphStyle(name, fontName=font, fontSize=size, leading=lead,
                              spaceBefore=sb, spaceAfter=sa, textColor=colors.HexColor(color))

    S = {
        "h1": st("h1", _FONT_B, 18, 22, "#15171d", 0, 1),
        "sub": st("sub", _FONT_B, 13, 17, "#3f8cff", 0, 1),
        "h2": st("h2", _FONT_B, 13, 17, "#17191f", 12, 6),
        "h3": st("h3", _FONT_B, 11, 15, "#1f232b", 9, 3),
        "body": st("body", _FONT, 10, 14, "#2b2f37"),
        "muted": st("muted", _FONT, 9, 12, "#7b808a"),
    }

    story = [
        Paragraph("Бизнес-отчёт: гипотезы НИОКР", S["h1"]),
        Paragraph(_p(project.get("title", "")), S["sub"]),
        Paragraph(f"Сформировано: {datetime.now().strftime('%d.%m.%Y %H:%M')}", S["muted"]),
        Paragraph("Цель проекта (KPI)", S["h2"]),
        Paragraph(_p(project.get("kpi_target", "")), S["body"]),
    ]
    if _kpi_line(project):
        story.append(Paragraph(_p(_kpi_line(project)), S["muted"]))
    if project.get("constraints"):
        story.append(Paragraph("Ограничения: " + _p(project["constraints"]), S["muted"]))

    s = _summary(hyps)
    story.append(Paragraph("Сводка и достижение KPI", S["h2"]))
    story.append(Paragraph(
        f"Всего гипотез: <b>{s['total']}</b> · подтверждено экспертом: <b>{s['confirmed']}</b> · "
        f"отклонено: <b>{s['refuted']}</b> · средний ранг: <b>{s['avg_composite']}</b>", S["body"]))
    story.append(Paragraph(
        f"Прогноз достижения KPI: <b>{s['high_value_count']}</b> гипотез с высоким ожидаемым "
        f"эффектом (ценность ≥ 70).", S["body"]))
    if s["top_statement"]:
        story.append(Paragraph(f"Лидер ранжирования ({s['top_composite']}): {_p(s['top_statement'])}", S["muted"]))

    story.append(Paragraph("Ранжированные гипотезы", S["h2"]))
    for i, h in enumerate(hyps, 1):
        sc = _scores(h)
        story.append(Paragraph(f"{i}. {_p(h.get('statement', ''))}", S["h3"]))
        story.append(Paragraph(
            f"Ранг: <b>{h.get('composite')}</b> · Статус: {_p(STATUS_LABELS.get(h.get('status'), h.get('status')))}",
            S["muted"]))
        for key, label in (("goal_link", "Связь с целью"), ("mechanism", "Механизм влияния"),
                           ("validation", "План проверки"), ("expert_notes", "Отзыв эксперта")):
            if h.get(key):
                story.append(Paragraph(f"<b>{label}:</b> {_p(h[key])}", S["body"]))
        row_l = [Paragraph(f"<font size=8 color='#7b808a'>{lbl}</font>", S["body"]) for _, lbl in DIMS]
        row_v = [Paragraph(f"<b>{_num(sc.get(k))}</b>", S["body"]) for k, _ in DIMS]
        tbl = Table([row_l, row_v], colWidths=[43 * mm] * 4)
        tbl.setStyle(TableStyle([
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#eceef2")),
            ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#eceef2")),
            ("TOPPADDING", (0, 0), (-1, -1), 4), ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        story.append(Spacer(1, 3))
        story.append(tbl)
        srcs = [e.get("title") for e in (h.get("evidence") or []) if e.get("title")]
        if srcs:
            story.append(Paragraph("Источники: " + _p("; ".join(srcs)), S["muted"]))
        story.append(Spacer(1, 6))

    buf = io.BytesIO()
    SimpleDocTemplate(
        buf, pagesize=A4, title="Бизнес-отчёт: гипотезы НИОКР",
        leftMargin=18 * mm, rightMargin=18 * mm, topMargin=16 * mm, bottomMargin=16 * mm,
    ).build(story)
    return buf.getvalue()


# ── DOCX (python-docx) ───────────────────────────────────────────────────────
def to_docx(project, hyps, weights):
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    doc.add_heading("Бизнес-отчёт: гипотезы НИОКР", level=0)
    doc.add_heading(project.get("title", ""), level=1)
    doc.add_paragraph(f"Сформировано: {datetime.now().strftime('%d.%m.%Y %H:%M')}")

    doc.add_heading("Цель проекта (KPI)", level=2)
    doc.add_paragraph(project.get("kpi_target", ""))
    if _kpi_line(project):
        doc.add_paragraph(_kpi_line(project))
    if project.get("constraints"):
        doc.add_paragraph("Ограничения: " + project["constraints"])

    s = _summary(hyps)
    doc.add_heading("Сводка и достижение KPI", level=2)
    doc.add_paragraph(
        f"Всего гипотез: {s['total']} · подтверждено экспертом: {s['confirmed']} · "
        f"отклонено: {s['refuted']} · средний ранг: {s['avg_composite']}")
    doc.add_paragraph(
        f"Прогноз достижения KPI: {s['high_value_count']} гипотез с высоким ожидаемым эффектом (ценность ≥ 70).")
    if s["top_statement"]:
        doc.add_paragraph(f"Лидер ранжирования ({s['top_composite']}): {s['top_statement']}")

    doc.add_heading("Ранжированные гипотезы", level=2)
    for i, h in enumerate(hyps, 1):
        sc = _scores(h)
        doc.add_heading(f"{i}. {h.get('statement', '')}", level=3)
        meta = doc.add_paragraph()
        run = meta.add_run(f"Ранг: {h.get('composite')} · Статус: "
                           f"{STATUS_LABELS.get(h.get('status'), h.get('status'))}")
        run.bold = True
        for key, label in (("goal_link", "Связь с целью"), ("mechanism", "Механизм влияния"),
                           ("validation", "План проверки"), ("expert_notes", "Отзыв эксперта")):
            if h.get(key):
                p = doc.add_paragraph()
                p.add_run(f"{label}: ").bold = True
                p.add_run(str(h[key]))
        table = doc.add_table(rows=2, cols=4)
        table.style = "Light Grid Accent 1"
        for col, (k, label) in enumerate(DIMS):
            table.cell(0, col).text = label
            table.cell(1, col).text = str(_num(sc.get(k)))
        srcs = [e.get("title") for e in (h.get("evidence") or []) if e.get("title")]
        if srcs:
            doc.add_paragraph("Источники: " + "; ".join(srcs))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

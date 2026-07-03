from ingestion.extractors.base import BaseExtractor, ExtractorError, missing_dependency
from ingestion.metadata import (
    build_base_metadata,
    clean_text,
    enrich_metadata,
    normalize_cell_value,
)
from ingestion.schemas import ParsedSection, ParsedTable, ParseResult


MAX_TABLE_ROWS = 5000


class DOCXExtractor(BaseExtractor):
    file_type = "docx"

    def extract(self, path: str, filename: str) -> ParseResult:
        try:
            from docx import Document
        except ImportError as exc:
            raise missing_dependency("python-docx", self.file_type) from exc

        warnings = []
        try:
            document = Document(path)
        except Exception as exc:
            raise ExtractorError(f"Failed to read DOCX: {exc}") from exc

        document_properties = self._core_properties(document)
        sections = self._sections(document)
        raw_text = "\n\n".join(section.text for section in sections if section.text)
        tables = self._tables(document, warnings)

        if not raw_text and not tables:
            warnings.append("DOCX contains no extractable text or tables")

        metadata = build_base_metadata(filename, self.file_type, document_properties=document_properties)
        metadata = enrich_metadata(
            metadata,
            text=raw_text,
            title=document_properties.get("title"),
            authors=document_properties.get("author"),
            keywords=document_properties.get("keywords"),
            year_values=(document_properties.get("created"), document_properties.get("modified")),
        )

        return ParseResult(
            filename=filename,
            file_type=self.file_type,
            raw_text=raw_text,
            sections=sections,
            tables=tables,
            metadata=metadata,
            warnings=warnings,
        )

    def _core_properties(self, document) -> dict:
        props = document.core_properties
        return {
            "title": props.title or None,
            "author": props.author or None,
            "subject": props.subject or None,
            "keywords": props.keywords or None,
            "comments": props.comments or None,
            "category": props.category or None,
            "created": props.created,
            "modified": props.modified,
            "last_modified_by": props.last_modified_by or None,
        }

    def _sections(self, document):
        sections = []
        title = "Document"
        buffer = []
        order = 1

        def flush():
            nonlocal order, buffer
            text = clean_text("\n".join(buffer))
            if text:
                sections.append(ParsedSection(title=title, text=text, order=order))
                order += 1
            buffer = []

        for paragraph in document.paragraphs:
            text = clean_text(paragraph.text)
            if not text:
                continue
            style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
            if style_name.startswith("heading"):
                flush()
                title = text
            else:
                buffer.append(text)
        flush()
        return sections

    def _tables(self, document, warnings):
        parsed = []
        for index, table in enumerate(document.tables, start=1):
            rows = []
            for row in table.rows:
                values = [normalize_cell_value(cell.text) for cell in row.cells]
                if any(value is not None for value in values):
                    rows.append(values)
            if not rows:
                continue

            headers = self._headers(rows[0])
            data_rows = rows[1:]
            if len(data_rows) > MAX_TABLE_ROWS:
                warnings.append(f"DOCX table {index} was truncated to {MAX_TABLE_ROWS} rows")
                data_rows = data_rows[:MAX_TABLE_ROWS]

            parsed.append(ParsedTable(
                name=f"Table {index}",
                sheet_name=None,
                columns=headers,
                rows=[self._row_dict(headers, row) for row in data_rows],
                data_range=f"R1C1:R{len(rows)}C{len(headers)}",
            ))
        return parsed

    def _headers(self, values):
        headers = []
        seen = {}
        for index, value in enumerate(values, start=1):
            header = str(value or f"Column {index}").strip() or f"Column {index}"
            count = seen.get(header.lower(), 0) + 1
            seen[header.lower()] = count
            headers.append(header if count == 1 else f"{header} {count}")
        return headers

    def _row_dict(self, headers, row):
        return {
            header: normalize_cell_value(row[index]) if index < len(row) else None
            for index, header in enumerate(headers)
        }

